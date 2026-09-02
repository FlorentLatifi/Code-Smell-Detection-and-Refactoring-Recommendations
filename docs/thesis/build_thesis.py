"""Generates the thesis document with the formatting UBT requires.

The formatting rules are not negotiable and are easy to break by hand: three
page-numbering regimes, a font rule per heading level, a line-spacing minimum.
Encoding them here means the document cannot drift out of compliance: rebuild
and it is correct again.

    python docs/thesis/build_thesis.py

Content lives in CONTENT below, so prose can be edited without touching layout
code. After the skeleton is handed over, further editing happens in Word;
rerunning this script overwrites the file.
"""

from __future__ import annotations

import os
from dataclasses import dataclass

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

FONT = "Times New Roman"
BODY_SIZE = Pt(12)
TITLE_SIZE = Pt(14)
CAPTION_SIZE = Pt(11)
LINE_SPACING = 1.5

from chapters import (  # noqa: E402
    CHAPTER_2,
    CHAPTER_3,
    CHAPTER_4,
    CHAPTER_6,
    abstract,
    chapter_5,
    chapter_8,
)
from references import all_references  # noqa: E402

OUTPUT = os.path.join(os.path.dirname(__file__), "Punim_Diplome_Florent_Latifi.docx")
LOGO = os.path.join(os.path.dirname(__file__), "assets", "ubt_logo.jpg")

# Anything the author still has to supply is marked so it cannot be missed
# during a final read-through.
TODO = "[PLOTËSO]"


def _set_properties(doc: Document) -> None:
    """Fill in the document properties Word shows under File > Info.

    python-docx starts from a blank template whose author is the library
    itself, and Word surfaces that string in the properties pane and in a
    printed footer if one is ever added.
    """
    properties = doc.core_properties
    properties.author = AUTHOR
    properties.last_modified_by = AUTHOR
    properties.title = TITLE_SQ
    properties.subject = "Punim diplome, Universiteti për Biznes dhe Teknologji"
    properties.comments = ""


# ----------------------------------------------------------------------
# Low-level Word plumbing
# ----------------------------------------------------------------------
def _set_font(run, size=BODY_SIZE, bold=False, caps=False):
    run.font.name = FONT
    run.font.size = size
    run.font.bold = bold
    run.font.all_caps = caps
    # Word keeps a separate east-Asian font slot; without this the theme font
    # can leak back in and the document silently stops being Times New Roman.
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)


def _field(paragraph, instruction: str, placeholder: str = "") -> None:
    """Insert a Word field (PAGE, TOC, ...) that Word evaluates on open."""
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = placeholder
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, instr, separate, text, end):
        run._r.append(element)
    _set_font(run)


def _page_numbering(section, fmt: str | None, start: int | None) -> None:
    """Set the numbering regime of one section.

    The template asks for three: none on the cover pages, roman for the front
    matter, arabic restarting at 1 for the chapters.
    """
    sect_pr = section._sectPr
    pg_num = sect_pr.find(qn("w:pgNumType"))
    if pg_num is None:
        pg_num = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num)
    if fmt:
        pg_num.set(qn("w:fmt"), fmt)
    if start is not None:
        pg_num.set(qn("w:start"), str(start))


def _footer_page_number(section, enabled: bool) -> None:
    """Page number bottom right, as the template illustrates."""
    section.footer.is_linked_to_previous = False
    paragraph = section.footer.paragraphs[0]
    paragraph.text = ""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.line_spacing = 1.0
    if enabled:
        _field(paragraph, " PAGE ", "1")


# ----------------------------------------------------------------------
# Building blocks
# ----------------------------------------------------------------------
def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = BODY_SIZE
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    fmt = normal.paragraph_format
    fmt.line_spacing = LINE_SPACING
    fmt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fmt.space_after = Pt(6)

    # Heading styles carry the required look *and* feed the table of contents,
    # so the TOC field can build itself instead of being typed by hand.
    for name, size, caps in (("Heading 1", TITLE_SIZE, True), ("Heading 2", BODY_SIZE, False)):
        style = doc.styles[name]
        style.font.name = FONT
        style.font.size = size
        style.font.bold = True
        style.font.all_caps = caps
        style.font.color.rgb = None
        style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        style.paragraph_format.line_spacing = LINE_SPACING
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.keep_with_next = True


def centered(doc: Document, text: str, *, size=BODY_SIZE, bold=False, caps=False, space_after=0):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.line_spacing = LINE_SPACING
    paragraph.paragraph_format.space_after = Pt(space_after)
    _set_font(paragraph.add_run(text), size=size, bold=bold, caps=caps)
    return paragraph


def blank(doc: Document, count: int = 1) -> None:
    for _ in range(count):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.line_spacing = 1.0
        paragraph.paragraph_format.space_after = Pt(0)


def body(doc: Document, text: str) -> None:
    """One body paragraph, with **emphasis** rendered rather than printed.

    The chapters were written with Markdown-style emphasis and it used to reach
    the page as literal asterisks: four paragraphs of the built document carried
    them, and nothing caught it, because the format check reads fonts and sizes
    rather than punctuation. Splitting on the marker keeps the emphasis the
    author meant and keeps the asterisks off the page; `check_format` fails the
    build if any survive.
    """
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing = LINE_SPACING
    paragraph.paragraph_format.first_line_indent = Pt(18)
    # Odd pieces are plain, even pieces are emphasised: "a **b** c" splits into
    # ["a ", "b", " c"]. An unmatched marker leaves one piece and stays literal,
    # which the check then reports rather than hiding.
    for index, piece in enumerate(text.split("**")):
        if piece:
            _set_font(paragraph.add_run(piece), bold=index % 2 == 1)


def unnumbered_heading(doc: Document, text: str) -> None:
    """A front-matter heading: same look as a chapter title, but no number.

    Level 1 is used so the heading is still picked up by the table of contents,
    which the template requires for the figure, table and glossary lists.
    """
    heading = doc.add_heading(level=1)
    heading.paragraph_format.page_break_before = True
    _set_font(heading.add_run(text), size=TITLE_SIZE, bold=True, caps=True)


def chapter(doc: Document, number: int, text: str) -> None:
    heading = doc.add_heading(level=1)
    heading.paragraph_format.page_break_before = True
    _set_font(heading.add_run(f"{number} {text}"), size=TITLE_SIZE, bold=True, caps=True)


def section_heading(doc: Document, number: str, text: str) -> None:
    heading = doc.add_heading(level=2)
    _set_font(heading.add_run(f"{number} {text}"), size=BODY_SIZE, bold=True)


def bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing = LINE_SPACING
    _set_font(paragraph.add_run(text))


@dataclass
class Numbering:
    """Numrat e figurave dhe të tabelave, të caktuar sipas radhës së shfaqjes.

    Shablloni e kërkon numërimin sipas radhës në të cilën lexuesi i takon, dhe ajo
    radhë është veti e dokumentit, jo e kapitullit ku u shkrua elementi. Numrat e
    shtypur me dorë e kishin humbur atë veti pa u vënë re: dokumenti hapej me
    «Figura 5». I caktuar këtu, numri s'ka si të mos përputhet me radhën.
    """

    figure: int = 0
    table: int = 0


def figure(doc: Document, path: str, text: str, numbering: Numbering) -> None:
    """Një figurë me përshkrimin poshtë saj, siç e kërkon shablloni."""
    holder = doc.add_paragraph()
    holder.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if os.path.exists(path):
        holder.add_run().add_picture(path, width=Inches(5.6))
    else:
        _set_font(holder.add_run(f"{TODO}: mungon figura {os.path.basename(path)}"))
    numbering.figure += 1
    caption(doc, f"Figura {numbering.figure}. {text}")


def table(
    doc: Document, text: str, headers: list[str], rows: list[list[str]], numbering: Numbering
) -> None:
    """Një tabelë me titullin sipër, siç e kërkon shablloni."""
    numbering.table += 1
    caption(doc, f"Tabela {numbering.table}. {text}")
    grid = doc.add_table(rows=1, cols=len(headers))
    grid.style = "Table Grid"
    for cell, header in zip(grid.rows[0].cells, headers, strict=True):
        cell.paragraphs[0].paragraph_format.line_spacing = 1.0
        _set_font(cell.paragraphs[0].add_run(header), size=CAPTION_SIZE, bold=True)
    for values in rows:
        cells = grid.add_row().cells
        for cell, value in zip(cells, values, strict=True):
            cell.paragraphs[0].paragraph_format.line_spacing = 1.0
            _set_font(cell.paragraphs[0].add_run(value), size=CAPTION_SIZE)
    blank(doc)


def caption(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.line_spacing = 1.0
    _set_font(paragraph.add_run(text), size=CAPTION_SIZE)


# ----------------------------------------------------------------------
# Document sections
# ----------------------------------------------------------------------
def build_logo(doc: Document) -> None:
    """The university banner that heads both cover pages of the template."""
    holder = doc.add_paragraph()
    holder.alignment = WD_ALIGN_PARAGRAPH.CENTER
    holder.paragraph_format.space_after = Pt(12)
    if os.path.exists(LOGO):
        holder.add_run().add_picture(LOGO, width=Inches(5.9))
    else:
        _set_font(holder.add_run(f"{TODO}: mungon logoja e UBT-së"), size=TITLE_SIZE, bold=True)


def build_cover(doc: Document) -> None:
    """Page 1 of the template: logo, programme, title, degree, author, date."""
    build_logo(doc)
    centered(doc, "Programi për Shkenca Kompjuterike dhe Inxhinieri", bold=True)
    blank(doc, 8)
    centered(doc, TITLE_SQ, size=TITLE_SIZE, bold=True, caps=True)
    blank(doc, 2)
    centered(doc, "Shkalla Bachelor", bold=True)
    blank(doc, 8)
    centered(doc, AUTHOR, bold=True)
    blank(doc, 6)
    centered(doc, SUBMISSION_DATE)
    centered(doc, "Prishtinë")


def build_inner_page(doc: Document) -> None:
    """Page 2: adds the academic year, the supervisor and the degree statement."""
    doc.add_page_break()
    build_logo(doc)
    centered(doc, "Programi për Shkenca Kompjuterike dhe Inxhinieri", bold=True)
    blank(doc, 2)
    centered(doc, "Punim Diplome", bold=True)
    centered(doc, f"Viti akademik {ACADEMIC_YEAR}")
    blank(doc, 4)
    centered(doc, AUTHOR, bold=True)
    blank(doc, 2)
    centered(doc, TITLE_SQ, size=TITLE_SIZE, bold=True, caps=True)
    blank(doc, 3)
    centered(doc, f"Mentore: {SUPERVISOR}", bold=True)
    blank(doc, 5)
    centered(doc, SUBMISSION_DATE)
    blank(doc, 3)
    centered(
        doc,
        "Ky punim është përpiluar dhe dorëzuar në përmbushjen e kërkesave "
        "të pjesshme për Shkallën Bachelor",
    )


def build_front_matter(doc: Document, figures: list[str], tables: list[str]) -> None:
    unnumbered_heading(doc, "Abstrakt")
    for paragraph in abstract():
        body(doc, paragraph)
    blank(doc)
    body(doc, f"Fjalë kyçe: {KEYWORDS}")

    unnumbered_heading(doc, "Mirënjohje")
    body(doc, ACKNOWLEDGEMENTS)

    unnumbered_heading(doc, "Përmbajtja")
    toc = doc.add_paragraph()
    toc.paragraph_format.line_spacing = LINE_SPACING
    _field(
        toc,
        ' TOC \\o "1-3" \\h \\z \\u ',
        "Kliko me të djathtën këtu dhe zgjidh 'Update Field' për ta gjeneruar përmbajtjen.",
    )

    unnumbered_heading(doc, "Lista e figurave")
    for entry in figures:
        body(doc, entry)

    unnumbered_heading(doc, "Lista e tabelave")
    for entry in tables:
        body(doc, entry)

    unnumbered_heading(doc, "Fjalori i termave")
    for term in GLOSSARY:
        bullet(doc, term)


def render_sections(doc: Document, sections: list, numbering: Numbering) -> None:
    """Një kapitull, çfarëdo qofshin llojet e elementeve brenda tij."""
    for number, title, paragraphs in sections:
        if number:
            section_heading(doc, number, title)
        for item in paragraphs:
            if not isinstance(item, tuple):
                body(doc, item)
            elif item[0] == "bullet":
                bullet(doc, item[1])
            elif item[0] == "figure":
                figure(doc, item[1], item[2], numbering)
            elif item[0] == "table":
                table(doc, item[1], item[2], item[3], numbering)
            else:
                raise ValueError(f"element i panjohur: {item[0]}")


def caption_lists(chapters: list) -> tuple[list[str], list[str]]:
    """Lista e figurave dhe e tabelave, e ndërtuar nga vetë kapitujt.

    Ballina renderohet para kapitujve, ndaj listat duhen ditur para se elementet
    të numërohen. Ecja këtu është e njëjta që bën `render_sections`, dhe pikërisht
    kjo e bën listën të pamundur ta humbasë një tabelë: e vetmja mënyrë për ta
    shtuar një tabelë në dokument është ta shtosh te kapitulli, të cilin e lexon
    edhe kjo funksion. Përputhja verifikohet nga `check_format.py`.
    """
    figures: list[str] = []
    tables: list[str] = []
    for sections in chapters:
        for _, _, paragraphs in sections:
            for item in paragraphs:
                if not isinstance(item, tuple):
                    continue
                if item[0] == "figure":
                    figures.append(f"Figura {len(figures) + 1}. {item[2]}")
                elif item[0] == "table":
                    tables.append(f"Tabela {len(tables) + 1}. {item[1]}")
    return figures, tables


def build_introduction(doc: Document, numbering: Numbering) -> None:
    chapter(doc, 1, "Hyrje")
    render_sections(doc, INTRODUCTION, numbering)


def build_references(doc: Document) -> None:
    """Chapter 7, numbered and alphabetical, in the template's own format.

    Only cited sources belong here; anything read but not cited goes under
    Bibliografia. The list is rendered from references.py rather than typed into
    the document, so a citation added to the text and a citation added to the
    list cannot drift apart.
    """
    for number, reference in enumerate(all_references(), 1):
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.line_spacing = LINE_SPACING
        paragraph.paragraph_format.left_indent = Inches(0.5)
        paragraph.paragraph_format.first_line_indent = Inches(-0.5)
        _set_font(paragraph.add_run(f"[{number}]	{reference}"))


def build_remaining_chapters(doc: Document, chapters: dict, numbering: Numbering) -> None:
    """Every chapter after the introduction.

    Chapters 2, 3, 4 and 6 are prose and live in ``chapters.py``. Chapters 5 and 8
    are built at this moment from ``data/results/``, so a rebuilt experiment and a
    rebuilt document cannot disagree. Chapter 7 renders the reference list.
    """
    for number, title in REMAINING_CHAPTERS:
        chapter(doc, number, title)
        if number in chapters:
            render_sections(doc, chapters[number], numbering)
        elif title == "Referencat":
            build_references(doc)
        else:
            body(doc, f"{TODO}")


def build() -> str:
    # Kapitujt 5 dhe 8 lexojnë `data/results/` sa herë thirren, ndaj ndërtohen një
    # herë: ballina i numëron elementet e tyre para se ata të renderohen, dhe dy
    # lexime të veçanta do të ishin dy burime të vërtete për të njëjtin dokument.
    chapters = {
        1: INTRODUCTION,
        2: CHAPTER_2,
        3: CHAPTER_3,
        4: CHAPTER_4,
        5: chapter_5(),
        6: CHAPTER_6,
        8: chapter_8(),
    }
    figures, tables = caption_lists([chapters[n] for n in sorted(chapters)])
    numbering = Numbering()

    doc = Document()
    _set_properties(doc)
    configure_styles(doc)

    # --- Section 1: cover pages, unnumbered -----------------------------
    cover = doc.sections[0]
    _page_numbering(cover, None, None)
    _footer_page_number(cover, enabled=False)
    build_cover(doc)
    build_inner_page(doc)

    # --- Section 2: front matter, roman numerals ------------------------
    front = doc.add_section(WD_SECTION.NEW_PAGE)
    _page_numbering(front, "upperRoman", 1)
    _footer_page_number(front, enabled=True)
    build_front_matter(doc, figures, tables)

    # --- Section 3: chapters, arabic restarting at 1 --------------------
    main = doc.add_section(WD_SECTION.NEW_PAGE)
    _page_numbering(main, "decimal", 1)
    _footer_page_number(main, enabled=True)
    build_introduction(doc, numbering)
    build_remaining_chapters(doc, chapters, numbering)

    doc.save(OUTPUT)
    return OUTPUT


# ======================================================================
# CONTENT
# ======================================================================
TITLE_SQ = "Detektimi i code smells dhe rekomandimet për refaktorim"
AUTHOR = "Florent Latifi"
SUPERVISOR = "Altina Salihu"
ACADEMIC_YEAR = "2025 - 2026"
SUBMISSION_DATE = "Nëntor / 2026"
KEYWORDS = (
    "code smells, refaktorim, metrika softuerike, cilësia e kodit, "
    "mësim i makinës, analizë statike, Java"
)

ACKNOWLEDGEMENTS = (
    f"{TODO}: Falënderimet shkruhen me fjalët e tua. Zakonisht përfshihen mentorja, "
    "profesorët e programit, familja dhe kushdo që ka kontribuar në këtë punim."
)

GLOSSARY = [
    "AST - Abstract Syntax Tree, pema e sintaksës abstrakte",
    "ATFD - Access To Foreign Data, qasja në të dhënat e klasave të tjera",
    "CBO - Coupling Between Objects, lidhshmëria mes objekteve",
    "CC - Cyclomatic Complexity, kompleksiteti ciklomatik",
    "CK - Chidamber & Kemerer, suita e metrikave të objekteve",
    "DIT - Depth of Inheritance Tree, thellësia e pemës së trashëgimisë",
    "FDP - Foreign Data Providers, numri i klasave që ofrojnë të dhëna të huaja",
    "LAA - Locality of Attribute Accesses, lokaliteti i qasjeve në atribute",
    "LCOM - Lack of Cohesion in Methods, mungesa e kohezionit në metoda",
    "LOC - Lines of Code, rreshtat e kodit",
    "MLCQ - Machine Learning Code Quality, dataset i etiketuar i code smells",
    "NOC - Number of Children, numri i nënklasave të drejtpërdrejta",
    "RFC - Response For a Class, përgjigjja e një klase",
    "TCC - Tight Class Cohesion, kohezioni i ngushtë i klasës",
    "WMC - Weighted Methods per Class, metodat e peshuara për klasë",
    "WOC - Weight of Class, pesha funksionale e klasës",
    # Metrikat e nivelit të metodës dhe termat e vlerësimit u shtuan pasi u
    # ndërtuan; fjalori duhet të mbulojë çdo shkurtesë që del në Kapitullin 5.
    "AMW - Average Method Weight, pesha mesatare e metodave",
    "CINT - Coupling Intensity, intensiteti i lidhshmërisë së një metode",
    "CLOC - Class Lines of Code, rreshtat efektivë të një klase",
    "F1 - mesatarja harmonike e precizionit dhe recall-it",
    "GroupKFold - ndarje me grupe, ku çdo depo bie e tëra në një fold",
    "MAXCC - kompleksiteti maksimal ciklomatik në një klasë",
    "MAXNESTING - thellësia maksimale e ndërfutjes së blloqeve",
    "MCC - Matthews Correlation Coefficient, koeficienti i korrelacionit i Matthews-it",
    "MLOC - Method Lines of Code, rreshtat efektivë të një metode",
    "NOAM - Number of Accessor Methods, numri i akses-metodave",
    "NOAV - Number of Accessed Variables, numri i variablave të prekura",
    "NOF - Number of Fields, numri i fushave",
    "NOM - Number of Methods, numri i metodave",
    "NOPA - Number of Public Attributes, numri i atributeve publike",
    "NP - Number of Parameters, numri i parametrave",
    "κ - kappa e Cohen-it, pajtimi mes dy vlerësuesve përtej rastësisë",
]

INTRODUCTION = [
    (
        None,
        "",
        [
            "Zhvillimi i një sistemi softuerik nuk përfundon me lëshimin e versionit "
            "të parë. Përkundrazi, pjesa më e madhe e jetës së një sistemi kalon në "
            "fazën e mirëmbajtjes dhe të evoluimit, ku kodi lexohet, kuptohet dhe "
            "ndryshohet vazhdimisht nga zhvillues të ndryshëm. Lehman (1980) e "
            "formuloi këtë si ligj: një sistem që përdoret në një mjedis real duhet "
            "të ndryshojë vazhdimisht, përndryshe bëhet gradualisht më pak i "
            "dobishëm. Rrjedhimisht, lehtësia me të cilën kodi mund të ndryshohet "
            "nuk është çështje estetike, por faktor ekonomik.",
            "Ky punim trajton një aspekt konkret të asaj lehtësie: strukturën e "
            "brendshme të kodit burimor dhe mënyrat automatike për ta vlerësuar e "
            "përmirësuar atë.",
        ],
    ),
    (
        "1.1",
        "Konteksti",
        [
            "Termi code smell u prezantua nga Kent Beck dhe u popullarizua nga "
            "Fowler (2018) në veprën Refactoring. Një code smell nuk është gabim: "
            "programi kompilohet, testet kalojnë dhe funksionaliteti është i saktë. "
            "Ai është një simptomë sipërfaqësore që zakonisht tregon një problem më "
            "të thellë të dizajnit. Një klasë që mban shumë përgjegjësi të palidhura "
            "mes tyre, një metodë që zgjatet për qindra rreshta, ose një klasë që "
            "mban vetëm të dhëna pa asnjë sjellje, janë shembuj tipikë.",
            "Rëndësia e tyre qëndron në efektin kumulativ. Cunningham (1992) e "
            "përshkroi këtë fenomen me metaforën e borxhit teknik: çdo kompromis i "
            "vogël në strukturë krijon një detyrim që paguhet me interes në çdo "
            "ndryshim të ardhshëm. Një sistem me borxh të lartë teknik nuk dështon "
            "menjëherë, por bëhet gjithnjë e më i shtrenjtë për t'u ndryshuar.",
            "Kundërpesha ndaj këtij degradimi është refaktorimi, i përkufizuar nga "
            "Fowler (2018) si ndryshim i strukturës së brendshme të kodit pa "
            "ndryshuar sjelljen e tij të jashtme. Refaktorimi presupozon dy gjëra: "
            "që problemi të jetë identifikuar, dhe që të dihet cila transformatë e "
            "adreson atë. Të dyja këto janë pikërisht objekt i këtij punimi.",
        ],
    ),
    (
        "1.2",
        "Motivimi",
        [
            "Identifikimi manual i code smells është i realizueshëm vetëm në shkallë "
            "të vogël. Në një sistem me qindra klasa, shqyrtimi i secilës prej tyre "
            "nga një zhvillues është praktikisht i pamundur dhe, për më tepër, i "
            "varur nga përvoja subjektive e shqyrtuesit.",
            "Mjetet ekzistuese të analizës statike e automatizojnë pjesërisht këtë "
            "punë, por vuajnë nga tri kufizime të dukshme. Së pari, shumica e tyre "
            "mbështeten në pragje fikse mbi një metrikë të vetme, prandaj një metodë "
            "e gjatë por koherente shënohet njësoj si një metodë e shkurtër që "
            "përzien disa përgjegjësi. Së dyti, ato fokusohen kryesisht në probleme "
            "të nivelit të rreshtit dhe të stilit, ndërsa problemet strukturore të "
            "nivelit të dizajnit mbulohen dobët. Së treti, dhe më e rëndësishmja, "
            "ato ndalen te njoftimi: i tregojnë zhvilluesit se çfarë është e "
            "gabuar, por jo se çfarë duhet bërë konkretisht.",
            "Kjo hapësirë mes identifikimit dhe veprimit është motivimi kryesor i "
            "këtij punimi.",
        ],
    ),
    (
        "1.3",
        "Qëllimi dhe objektivat",
        [
            "Qëllimi i këtij punimi është projektimi, implementimi dhe vlerësimi "
            "empirik i një sistemi që identifikon code smells në kod burimor Java "
            "dhe propozon refaktorime konkrete e të verifikueshme për t'i adresuar "
            "ato.",
            "Për ta arritur këtë qëllim janë përcaktuar objektivat e mëposhtme:",
            ("bullet", "Të shqyrtohet literatura mbi metrikat e cilësisë së kodit, "
             "strategjitë e detektimit të code smells dhe teknikat e refaktorimit."),
            ("bullet", "Të implementohet një motor analize që nxjerr metrika të "
             "matshme nga kodi burimor Java në nivel klase dhe metode."),
            ("bullet", "Të implementohet detektimi i bazuar në rregulla, duke "
             "përdorur strategji të publikuara dhe pragje të justifikuara nga "
             "literatura."),
            ("bullet", "Të trajnohet dhe vlerësohet një model i mësimit të makinës "
             "mbi një dataset të etiketuar nga zhvillues profesionistë, për ta "
             "krahasuar me qasjen e bazuar në rregulla."),
            ("bullet", "Të implementohet një motor refaktorimi që gjeneron "
             "transformime konkrete dhe verifikon objektivisht efektin e tyre."),
            ("bullet", "Të ndërtohet një ndërfaqe web që i bën rezultatet të "
             "shfrytëzueshme nga zhvilluesi."),
        ],
    ),
    (
        "1.4",
        "Pyetjet kërkimore",
        [
            "Punimi synon t'u përgjigjet tri pyetjeve kërkimore, secila e matshme "
            "me kritere objektive:",
            ("bullet", "PK1: Sa e saktë është detektimi i bazuar në strategji "
             "metrikash, krahasuar me etiketimet manuale të zhvilluesve "
             "profesionistë?"),
            ("bullet", "PK2: A e përmirëson një model i mësimit të makinës, i "
             "trajnuar mbi të njëjtat metrika, saktësinë e detektimit krahasuar me "
             "pragjet fikse?"),
            ("bullet", "PK3: A i përmirësojnë objektivisht refaktorimet e propozuara "
             "karakteristikat strukturore të kodit, duke ruajtur "
             "kompilueshmërinë dhe sjelljen e tij?"),
        ],
    ),
    (
        "1.5",
        "Fushëveprimi dhe kufizimet",
        [
            "Punimi kufizohet në gjuhën programuese Java. Kjo zgjedhje është bërë "
            "sepse pjesa dërrmuese e literaturës mbi metrikat e objekteve dhe "
            "datasetet e etiketuara të code smells janë ndërtuar mbi kod Java, çka "
            "mundëson krahasim të drejtpërdrejtë me rezultatet e publikuara.",
            "Analiza kryhet mbi kodin burimor në mënyrë statike, pa ekzekutim të "
            "programit. Rrjedhimisht, karakteristikat që shfaqen vetëm gjatë "
            "ekzekutimit nuk mbulohen. Po ashtu, sistemi nuk kryen zgjidhje të "
            "plotë të tipave, prandaj disa varësi që kërkojnë analizë të thellë "
            "semantike trajtohen në mënyrë konservative.",
            "Grupi i code smells të mbuluara është i kufizuar te ata për të cilët "
            "ekzistojnë strategji detektimi të publikuara dhe të dhëna të "
            "etiketuara, çka mundëson vlerësim empirik të besueshëm.",
            f"{TODO}: Shto këtu çdo kufizim tjetër që dalin gjatë punës, "
            "për shembull madhësia e korpusit ose kufizimet e verifikimit.",
        ],
    ),
    (
        "1.6",
        "Struktura e punimit",
        [
            "Punimi është organizuar si vijon. Kapitulli 2 shqyrton literaturën mbi "
            "metrikat e cilësisë, strategjitë e detektimit dhe mjetet ekzistuese. "
            "Kapitulli 3 formulon problemin që adresohet. Kapitulli 4 përshkruan "
            "metodologjinë, arkitekturën e sistemit dhe vendimet e projektimit. "
            "Kapitulli 5 paraqet rezultatet e vlerësimit empirik. Kapitulli 6 "
            "diskuton gjetjet, kufizimet dhe drejtimet e punës së ardhshme. "
            "Kapitulli 7 përmban referencat dhe kapitulli 8 shtojcat.",
        ],
    ),
]

REMAINING_CHAPTERS = [
    (2, "Shqyrtimi i literaturës"),
    (3, "Deklarimi i problemit"),
    (4, "Metodologjia"),
    (5, "Rezultatet"),
    (6, "Diskutime dhe përfundime"),
    (7, "Referencat"),
    (8, "Shtojcat"),
]


if __name__ == "__main__":
    print(f"U gjenerua: {build()}")

    # Importuar këtu e jo lart, sepse check_citations e lexon këtë modul: në krye
    # do të ishte import qarkor. Ndërtimi nuk ndalet nga një mospërputhje — gjatë
    # shkrimit ajo është gjendje kalimtare — por nuk kalon as pa u thënë.
    from check_citations import report

    problems = report()
    if problems:
        print("Kujdes, teksti dhe lista e referencave nuk përputhen:")
        print("\n".join(problems))
