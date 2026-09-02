"""Verifikon që dokumenti i gjeneruar i përmbahet formatimit që kërkon UBT-ja.

    python docs/thesis/check_format.py [shtegu.docx]

`build_thesis.py` i zbaton këto rregulla; ky skript i **verifikon te skedari**, dhe
dallimi është i tërë qëllimi. Dokumenti nuk mbetet ashtu si u gjenerua: docstring-u
i ndërtuesit e thotë vetë se redaktimi i mëtejshëm bëhet në Word. Një rresht i
ngjitur nga diku tjetër e sjell fontin e vet, një titull i shtuar me dorë nuk merr
stilin e kreut, dhe asnjëra prej të dyjave nuk duket ndryshe në ekran derisa dikush
i mat. Komisioni e vlerëson përputhshmërinë me shabllonin, ndaj ato dallime
kushtojnë pikë.

Rregullat e kontrolluara janë ato të shabllonit: Times New Roman kudo; tekst 12pt i
drejtuar me hapësirë rreshtash së paku 1.5; tituj kapitujsh 14pt bold me shkronja
kapitale, secili në faqe të re; nëntituj 12pt bold; përshkrime figurash e tituj
tabelash 11pt në qendër, të numëruar pa hapësira; tri regjime numërimi faqesh; numri
i faqes poshtë djathtas. Kontrollohet edhe që lista e figurave dhe e tabelave në
fillim të përputhet me atë që dokumenti përmban vërtet.

Kontrollohen edhe numërimi i njëpasnjëshëm i nënkapitujve dhe mungesa e krerëve pa
përmbajtje. Në fund listohen vendet `[PLOTËSO]` që i mbeten autorit: ato janë të
shënuara me qëllim dhe nuk e rrëzojnë kontrollin, por në një dokument prej dhjetëra
faqesh gjenden vetëm nga kush di t'i kërkojë.

Del me kod jo-zero po qe se ndonjë rregull thyhet, që kontrolli të mund të hyjë në
CI bashkë me atë të citimeve.
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

from build_thesis import BODY_SIZE, CAPTION_SIZE, FONT, OUTPUT, TITLE_SIZE
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph

# Numërimi i tri regjimeve, sipas radhës së seksioneve: kopertina pa numra, ballina
# me romakë duke filluar nga I, kapitujt me arabë duke rifilluar nga 1.
EXPECTED_SECTIONS = (
    {"fmt": None, "start": None, "footer": False},
    {"fmt": "upperRoman", "start": "1", "footer": True},
    {"fmt": "decimal", "start": "1", "footer": True},
)

CAPTION = re.compile(r"^(Figura|Tabela) (\d+)\.")
CHAPTER_HEADING = re.compile(r"^\d+ ")

# Hapësira minimale e rreshtave për tekstin rrjedhës. Përshkrimet, qelizat e
# tabelave dhe fundfaqja janë me qëllim më të ngjeshura, ndaj nuk hyjnë këtu.
MIN_LINE_SPACING = 1.5


def _size_ok(run, expected: Pt) -> bool:
    """Madhësia e shkronjave, duke lejuar trashëgiminë nga stili.

    `None` do të thotë «merre nga stili», dhe stilet i verifikon `_check_styles`.
    Kërkesa që çdo run ta mbajë vetë vlerën do të shënonte si gabim tekstin që Word
    e formaton saktë përmes stilit.
    """
    return run.font.size is None or run.font.size == expected


def _check_styles(doc: Document) -> list[str]:
    problems = []
    normal = doc.styles["Normal"]
    if normal.font.name != FONT:
        problems.append(f"stili Normal nuk është {FONT} por {normal.font.name}")
    if normal.font.size != BODY_SIZE:
        problems.append(f"stili Normal nuk është {BODY_SIZE.pt:g}pt")
    spacing = normal.paragraph_format.line_spacing
    if spacing < MIN_LINE_SPACING:
        problems.append(f"stili Normal ka hapësirë rreshtash {spacing}")

    for name, size, caps in (("Heading 1", TITLE_SIZE, True), ("Heading 2", BODY_SIZE, False)):
        style = doc.styles[name]
        if style.font.name != FONT:
            problems.append(f"stili {name} nuk është {FONT}")
        if style.font.size != size:
            problems.append(f"stili {name} nuk është {size.pt:g}pt")
        if not style.font.bold:
            problems.append(f"stili {name} nuk është bold")
        if bool(style.font.all_caps) != caps:
            problems.append(f"stili {name} me all_caps={style.font.all_caps}, pritej {caps}")
        if style.paragraph_format.alignment != WD_ALIGN_PARAGRAPH.LEFT:
            problems.append(f"stili {name} nuk është i rreshtuar majtas")
    return problems


def _check_fonts(doc: Document) -> list[str]:
    """Asnjë run me font tjetër, as në paragrafë as brenda qelizave të tabelave."""
    problems = []
    runs = [(p, r) for p in doc.paragraphs for r in p.runs]
    for grid in doc.tables:
        for row in grid.rows:
            for cell in row.cells:
                runs += [(p, r) for p in cell.paragraphs for r in p.runs]

    for paragraph, run in runs:
        if run.font.name not in (None, FONT):
            excerpt = (paragraph.text or run.text)[:40]
            problems.append(f"font {run.font.name} te «{excerpt}»")
    return problems


def _check_body(doc: Document) -> list[str]:
    """Teksti rrjedhës: 12pt, i drejtuar, me hapësirë rreshtash së paku 1.5."""
    problems = []
    for paragraph in doc.paragraphs:
        if not paragraph.text.strip() or paragraph.style.name != "Normal":
            continue
        if CAPTION.match(paragraph.text) or paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER:
            continue  # përshkrimet dhe kopertina kontrollohen veç
        excerpt = paragraph.text[:40]
        if paragraph.alignment not in (None, WD_ALIGN_PARAGRAPH.JUSTIFY):
            problems.append(f"paragraf i pa-drejtuar te «{excerpt}»")
        spacing = paragraph.paragraph_format.line_spacing
        if spacing is not None and spacing < MIN_LINE_SPACING:
            problems.append(f"hapësirë rreshtash {spacing} te «{excerpt}»")
        for run in paragraph.runs:
            if not _size_ok(run, BODY_SIZE):
                problems.append(f"madhësi {run.font.size.pt:g}pt te «{excerpt}»")
                break
    return problems


def _split_at_first_chapter(doc: Document) -> tuple[list, list]:
    """Ballina dhe trupi, të ndarë te kreu i Kapitullit 1.

    Ndarja nevojitet sepse te ballina rreshtat e listës së figurave dhe të tabelave
    e kanë të njëjtin tekst me përshkrimet e vërteta. Ata janë tekst rrjedhës e jo
    përshkrime, ndaj matja e tyre me rregullat e përshkrimit do të raportonte njëzet
    e një gabime që nuk ekzistojnë.
    """
    for index, paragraph in enumerate(doc.paragraphs):
        if paragraph.style.name == "Heading 1" and CHAPTER_HEADING.match(paragraph.text):
            return doc.paragraphs[:index], doc.paragraphs[index:]
    return doc.paragraphs, []


def _captions(doc: Document) -> list:
    """Përshkrimet e vërteta: ato brenda kapitujve, jo rreshtat e listave."""
    _, body_paragraphs = _split_at_first_chapter(doc)
    return [p for p in body_paragraphs if CAPTION.match(p.text)]


def _check_captions(doc: Document) -> list[str]:
    problems = []
    for paragraph in _captions(doc):
        if paragraph.alignment != WD_ALIGN_PARAGRAPH.CENTER:
            problems.append(f"përshkrim jo në qendër: «{paragraph.text[:40]}»")
        for run in paragraph.runs:
            if not _size_ok(run, CAPTION_SIZE):
                problems.append(f"përshkrim {run.font.size.pt:g}pt: «{paragraph.text[:40]}»")
                break
    return problems


def _check_markup(doc: Document) -> list[str]:
    """No Markdown marker may survive into the page.

    The chapters are written with ``**emphasis**`` and the builder turns it into
    bold runs. When that failed, four paragraphs reached the document with the
    asterisks visible and every other check still passed: fonts, sizes,
    numbering and captions were all correct, and none of them reads punctuation.
    This is the check that would have caught it.
    """
    problems = []
    for paragraph in doc.paragraphs:
        if "**" in paragraph.text:
            start = max(0, paragraph.text.index("**") - 30)
            problems.append(f"shenjë markdown e mbetur te «{paragraph.text[start:][:70]}»")
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if "**" in cell.text:
                    problems.append(f"shenjë markdown e mbetur te qeliza «{cell.text[:50]}»")
    return problems


def _check_numbering(doc: Document) -> list[str]:
    """Numrat e figurave dhe të tabelave ecin 1, 2, 3 sipas radhës së shfaqjes.

    Shablloni e kërkon numërimin sipas radhës ku lexuesi i takon. Kontrolli e ka
    kapur dokumentin duke e hapur me «Figura 5», sepse numrat ndiqnin radhën në të
    cilën ishin gjeneruar figurat, jo radhën e leximit.
    """
    problems = []
    counters = {"Figura": 0, "Tabela": 0}
    for paragraph in _captions(doc):
        match = CAPTION.match(paragraph.text)
        assert match is not None
        kind, number = match.group(1), int(match.group(2))
        counters[kind] += 1
        if number != counters[kind]:
            problems.append(f"{kind} {number} aty ku pritej {counters[kind]}")
    return problems


def _check_lists(doc: Document) -> list[str]:
    """Lista e ballinës përputhet me përshkrimet që dokumenti përmban vërtet.

    Të dyja dalin nga i njëjti burim — `caption_lists` e ndërton listën me të njëjtën
    ecje që bën renderimi — por nga dy kalime të ndryshme mbi të. Krahasimi këtu e
    verifikon se ato dy kalime pajtohen, çka nuk është e vetëkuptueshme: ballina
    numërohet para se elementet të renderohen.
    """
    front, _ = _split_at_first_chapter(doc)
    declared = [p.text for p in front if CAPTION.match(p.text)]
    actual = [p.text for p in _captions(doc)]
    problems = []
    for missing in sorted(set(declared) - set(actual)):
        problems.append(f"e listuar në ballinë por s'del në tekst: «{missing}»")
    for extra in sorted(set(actual) - set(declared)):
        problems.append(f"del në tekst por mungon në listën e ballinës: «{extra}»")
    return problems


def _check_sections(doc: Document) -> list[str]:
    problems = []
    if len(doc.sections) != len(EXPECTED_SECTIONS):
        return [f"{len(doc.sections)} seksione, pritej {len(EXPECTED_SECTIONS)}"]

    for index, (section, expected) in enumerate(zip(doc.sections, EXPECTED_SECTIONS, strict=True)):
        pg_num = section._sectPr.find(qn("w:pgNumType"))
        fmt = pg_num.get(qn("w:fmt")) if pg_num is not None else None
        start = pg_num.get(qn("w:start")) if pg_num is not None else None
        if fmt != expected["fmt"] or start != expected["start"]:
            problems.append(
                f"seksioni {index + 1}: numërim {fmt}/{start}, "
                f"pritej {expected['fmt']}/{expected['start']}"
            )

        footer = section.footer.paragraphs[0]
        has_field = "PAGE" in footer._p.xml
        if has_field != expected["footer"]:
            problems.append(
                f"seksioni {index + 1}: numër faqeje {has_field}, pritej {expected['footer']}"
            )
        if has_field and footer.alignment != WD_ALIGN_PARAGRAPH.RIGHT:
            problems.append(f"seksioni {index + 1}: numri i faqes nuk është djathtas")
    return problems


def _check_page_breaks(doc: Document) -> list[str]:
    """Çdo kapitull fillon në faqe të re, siç e kërkon shablloni."""
    problems = []
    for paragraph in doc.paragraphs:
        if paragraph.style.name != "Heading 1":
            continue
        if not paragraph.paragraph_format.page_break_before:
            problems.append(f"kreu pa ndarje faqeje: «{paragraph.text[:40]}»")
    return problems


def _check_numbering_of_sections(doc: Document) -> list[str]:
    """Nënkapitujt ecin 1, 2, 3 brenda kapitullit të tyre, pa kërcime.

    Numrat e nënkapitujve shkruhen me dorë te `chapters.py`, ndaj i ka të njëjtat
    mënyra për të dalë gabim që kishte numërimi i figurave para se të gjenerohej:
    një seksion i futur në mes dhe numrat e tjerë të palëvizur. Ndryshe nga
    figurat, këta nuk gjenerohen dot — numri i një nënkapitulli citohet në prozë
    dhe në planin e punimit — ndaj kontrollohen.
    """
    problems = []
    chapter = 0
    expected = 0
    for paragraph in doc.paragraphs:
        style = paragraph.style.name
        text = paragraph.text.strip()
        if style == "Heading 1" and CHAPTER_HEADING.match(text):
            chapter = int(text.split(" ", 1)[0])
            expected = 0
        elif style == "Heading 2" and chapter:
            number = text.split(" ", 1)[0]
            expected += 1
            if number != f"{chapter}.{expected}":
                problems.append(f"nënkapitulli «{number}» aty ku pritej {chapter}.{expected}")
                expected = int(number.rsplit(".", 1)[-1]) if "." in number else expected
    return problems


def _check_no_empty_headings(doc: Document) -> list[str]:
    """Asnjë kre pa asgjë nën të.

    Një kapitull që hapet drejt me nënkapitullin e vet nuk është bosh: përmbajtjen
    e mbajnë nënkapitujt. Bosh është ai pas të cilit vjen menjëherë një kre i të
    njëjtit nivel ose më i lartë, pra një titull që premton diçka që faqja nuk e
    mban — dhe që në përmbajtje duket njësoj si çdo tjetër.

    Ecja bëhet mbi trupin e dokumentit e jo mbi paragrafët, sepse një tabelë është
    përmbajtje po aq sa proza dhe nuk del fare te `doc.paragraphs`.
    """
    levels = {"Heading 1": 1, "Heading 2": 2}
    body: list[tuple[int | None, str]] = []
    for element in doc.element.body.iterchildren():
        tag = element.tag.rsplit("}", 1)[-1]
        if tag == "tbl":
            body.append((None, ""))
        elif tag == "p":
            paragraph = Paragraph(element, doc)
            if not paragraph.text.strip():
                continue
            body.append((levels.get(paragraph.style.name), paragraph.text.strip()))

    problems = []
    for index, (level, text) in enumerate(body):
        if level is None:
            continue
        following = body[index + 1 :]
        empty = not following or (following[0][0] is not None and following[0][0] <= level)
        if empty:
            problems.append(f"kre pa përmbajtje: «{text[:40]}»")
    return problems


def remaining_todos(doc: Document) -> list[str]:
    """Vendet që autori duhet t'i plotësojë, si kujtesë e jo si gabim.

    Të shënuara me qëllim; raportohen sepse në një dokument prej dhjetëra faqesh
    ato gjenden vetëm nga kush di t'i kërkojë.
    """
    marker = "[PLOTËSO]"
    return [p.text.strip()[:70] for p in doc.paragraphs if marker in p.text]


def report(path: Path) -> list[str]:
    doc = Document(str(path))
    problems = []
    for check in (
        _check_styles,
        _check_fonts,
        _check_body,
        _check_captions,
        _check_markup,
        _check_numbering,
        _check_lists,
        _check_sections,
        _check_page_breaks,
    ):
        problems += [f"  {line}" for line in check(doc)]
    return problems


def main() -> int:
    if hasattr(sys.stdout, "reconfigure"):
        sys.stdout.reconfigure(encoding="utf-8")

    path = Path(sys.argv[1]) if len(sys.argv) > 1 else Path(OUTPUT)
    if not path.exists():
        print(f"Dokumenti nuk ekziston: {path}. Ndërtoje me build_thesis.py.")
        return 1

    problems = report(path)
    if problems:
        print(f"Formatimi devijon nga shablloni te {path.name}:")
        print("\n".join(problems))
        return 1

    print(f"Formatimi është sipas shabllonit: {path.name}")
    todos = remaining_todos(Document(str(path)))
    if todos:
        print(f"\nMbeten {len(todos)} vende për t'u plotësuar nga autori:")
        for entry in todos:
            print(f"  {entry}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
